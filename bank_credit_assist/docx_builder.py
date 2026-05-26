"""
docx_builder.py
DocxReportBuilder — 从零构建对公信贷审查报告（python-docx）

完全舍弃 Word 模板，程序化构建文档。
数据来源：Phase2 结构数据 + Phase2.5 AI 推理文本。
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT

import pandas as pd
import re

from shared.utils import safe_print as _safe_print

# ── 中文字体配置 ────────────────────────────────────────────────
FONT_HEITI = "黑体"
FONT_BODY = "微软雅黑"
FONT_SIZE_COVER = Pt(22)
FONT_SIZE_H1 = Pt(16)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_BODY = Pt(10.5)  # 五号
FONT_SIZE_TABLE = Pt(9)
COLOR_DARK = RGBColor(0x00, 0x33, 0x66)
COLOR_BODY = RGBColor(0x00, 0x00, 0x00)
COLOR_MISSING = RGBColor(0xCC, 0x00, 0x00)
COLOR_GRAY = RGBColor(0x33, 0x33, 0x33)
MISSING_MARKER = "【待人工补充】"

_tz = timezone(timedelta(hours=8))


# ============================================================================
# ReportData — 统一数据容器
# ============================================================================

@dataclass
class ReportData:
    enterprise_name: str
    report_date: str = field(default_factory=lambda: datetime.now(_tz).strftime("%Y年%m月%d日"))
    basic_info: dict = field(default_factory=dict)
    financial_metrics: dict = field(default_factory=dict)
    tech_metrics: dict = field(default_factory=dict)
    inference_text: dict = field(default_factory=dict)
    income_verification: dict = field(default_factory=dict)
    guarantee_types: list[str] = field(default_factory=list)
    markdown_content: str = ""
    annual_indicators: dict = field(default_factory=dict)  # {year: {指标: 值}}


# ============================================================================
# DocxReportBuilder
# ============================================================================

class DocxReportBuilder:
    """从零构建对公信贷审查报告"""

    def __init__(self, data: ReportData) -> None:
        self.data = data
        self.doc = Document()
        self._tables_cache: dict | None = None

    # ── 页设置 ─────────────────────────────────────────────────

    def _setup_page(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ── 格式化辅助 ─────────────────────────────────────────────

    def _add_heading_1(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = FONT_HEITI
        run.font.size = FONT_SIZE_H1
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK
        self._set_east_asian_font(run, FONT_HEITI)

    def _add_heading_2(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = FONT_HEITI
        run.font.size = FONT_SIZE_H2
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK
        self._set_east_asian_font(run, FONT_HEITI)

    def _add_para(self, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
        if not text:
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(str(text))
        run.font.name = FONT_BODY
        run.font.size = FONT_SIZE_BODY
        run.font.bold = bold
        run.font.color.rgb = color or COLOR_BODY
        self._set_east_asian_font(run, FONT_BODY)

    def _add_kv(self, label: str, value: Any) -> None:
        """键值对行，如"法定代表人：张三" """
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        # 标签
        run_label = p.add_run(f"{label}：")
        run_label.font.name = FONT_BODY
        run_label.font.size = FONT_SIZE_BODY
        run_label.font.bold = True
        self._set_east_asian_font(run_label, FONT_BODY)
        # 值
        display = str(value) if value is not None and str(value).strip() else MISSING_MARKER
        run_val = p.add_run(display)
        run_val.font.name = FONT_BODY
        run_val.font.size = FONT_SIZE_BODY
        run_val.font.color.rgb = COLOR_MISSING if display == MISSING_MARKER else COLOR_BODY
        self._set_east_asian_font(run_val, FONT_BODY)

    def _add_missing(self, text: str = "") -> None:
        """添加缺失数据标注"""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text or MISSING_MARKER)
        run.font.name = FONT_BODY
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = COLOR_MISSING
        self._set_east_asian_font(run, FONT_BODY)

    def _get_text(self, key: str, default: str = "") -> str:
        """安全获取推理文本"""
        text = self.data.inference_text.get(key, "")
        if "【生成失败】" in str(text) or "【跳过】" in str(text):
            return default or MISSING_MARKER
        if not text or not str(text).strip():
            return "根据现有资料无法确定，建议实地调研补充"
        return str(text)

    def _get_value(self, key: str, data_dict: dict | None = None) -> Any:
        """安全获取数值"""
        d = data_dict or self.data.financial_metrics
        return d.get(key)

    @staticmethod
    def _set_east_asian_font(run, font_name: str) -> None:
        """设置东亚文字字体"""
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = __import__('lxml').etree.SubElement(rPr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), font_name)

    def _add_table_caption(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(text)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = COLOR_GRAY
        self._set_east_asian_font(run, FONT_BODY)

    # ── 封面 ───────────────────────────────────────────────────

    def _add_cover_page(self) -> None:
        # 空行留白
        for _ in range(4):
            self.doc.add_paragraph()

        # 标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("公司客户综合授信调查报告")
        run.font.name = FONT_HEITI
        run.font.size = FONT_SIZE_COVER
        run.font.bold = True
        run.font.color.rgb = COLOR_DARK
        self._set_east_asian_font(run, FONT_HEITI)

        self.doc.add_paragraph()

        # 企业名称
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"企业名称：{self.data.enterprise_name}")
        run.font.name = FONT_BODY
        run.font.size = Pt(14)
        self._set_east_asian_font(run, FONT_BODY)

        # 报告日期
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"报告日期：{self.data.report_date}")
        run.font.name = FONT_BODY
        run.font.size = Pt(14)
        self._set_east_asian_font(run, FONT_BODY)

        # 风险等级
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"综合风险等级：{self.data.inference_text.get('risk_level', '待评估')}")
        run.font.name = FONT_BODY
        run.font.size = Pt(14)
        self._set_east_asian_font(run, FONT_BODY)

        self.doc.add_page_break()

    # ========================================================================
    # 各章构建
    # ========================================================================

    def _add_chapter_1(self) -> None:
        """一、申请人基本信息"""
        self._add_heading_1("一、申请人基本信息")

        # (一) 基本情况
        self._add_heading_2("（一）申请人基本情况")
        bi = self.data.basic_info
        self._add_kv("企业名称", bi.get("company_name") or self.data.enterprise_name)
        self._add_kv("统一社会信用代码", bi.get("unified_social_credit_code"))
        self._add_kv("法定代表人", bi.get("legal_representative"))
        self._add_kv("注册资本", bi.get("registered_capital"))
        reg_date = bi.get("registration_date") or bi.get("registrationDate") or bi.get("成立日期") or bi.get("注册日期")
        self._add_kv("注册日期", reg_date)
        self._add_kv("经营范围", bi.get("business_scope"))

        # (二) 股东结构
        self._add_heading_2("（二）股权结构及实际控制人")
        self._add_para(self._get_text("actual_controller"))

        # (三) 管理层信息
        self._add_heading_2("（三）管理层信息")
        self._add_para(self._get_text("management_team_summary"))

        # (四) 集团情况
        self._add_heading_2("（四）所属集团情况介绍")
        self._add_para(self._get_text("group_introduction"))

    def _add_chapter_2(self) -> None:
        """二、申请人经营情况"""
        self._add_heading_1("二、申请人经营情况")

        # (一) 经营概况
        self._add_heading_2("（一）经营概况")
        self._add_para(self._get_text("business_model"))
        self._add_para(self._get_text("operation_analysis"))
        self._add_kv("主营业务产品", self._get_text("main_products"))

        # (二) 核心技术
        self._add_heading_2("（二）核心技术及研发能力")
        self._add_para(self._get_text("core_technology"))
        tech = self.data.tech_metrics
        cert = tech.get("high_tech_cert", {}) if isinstance(tech.get("high_tech_cert"), dict) else {}
        self._add_kv("高新技术企业证书", cert.get("value", "") if cert else "")
        # 研发费用占比：tech_metrics 中为嵌套 dict {value: 11.56, ...} 或 financial_metrics 中为 flat 值
        rd_data = tech.get("rd_expense_ratio", {})
        rd_val = rd_data.get("value") if isinstance(rd_data, dict) else rd_data
        if rd_val is None:
            rd_val = self._get_value("rd_expense_ratio")
        self._add_kv("研发费用占比", rd_val)
        # 核心技术人员
        team_val = tech.get("team_size") or tech.get("core_team_size")
        if isinstance(team_val, dict):
            team_val = team_val.get("value")
        self._add_kv("核心技术人员", team_val)

        # (三) 产能产量
        self._add_heading_2("（三）产能产量情况")
        self._add_para(self._get_text("capacity_output_summary"))

        # (四) 毛利率分析
        self._add_heading_2("（四）毛利率分析")
        self._add_para(self._get_text("gross_margin_analysis"))

        # (五) 供应商
        self._add_heading_2("（五）主要供应商情况")
        self._add_para(self._get_text("upstream_suppliers"))

        # (六) 销售商
        self._add_heading_2("（六）主要销售商情况")
        self._add_para(self._get_text("downstream_customers"))

        # (七) 在手订单
        self._add_heading_2("（七）在手订单情况")
        self._add_para(self._get_text("current_orders_summary"))

        # (八) 在建项目
        self._add_heading_2("（八）在建项目及重大投资")
        self._add_para(self._get_text("major_investments"))

    def _add_chapter_3(self) -> None:
        """三、申请人财务状况"""
        self._add_heading_1("三、申请人财务状况")

        fm = self.data.financial_metrics
        is_consolidated = fm.get("is_consolidated", False)

        # (一) 财务报表
        self._add_heading_2("（一）财务报表")
        if is_consolidated:
            self._add_para("注：企业提供合并财务报表，以下分析基于合并报表数据。")
        else:
            self._add_para("注：企业提供单体财务报表。")

        # 尝试从 Phase1 markdown 提取三张报表
        bs_df = self._try_extract_table_from_md(["资产总计", "总资产", "流动资产", "负债"])
        is_df = self._try_extract_table_from_md(["营业收入", "净利润", "利润总额", "营业成本"])
        cf_df = self._try_extract_table_from_md(["经营活动", "投资活动", "筹资活动", "现金流量"])

        # 资产负债表
        if bs_df is not None:
            self._add_financial_table("资产负债表（单位：万元）", bs_df)
        else:
            self._add_table_caption("资产负债表（单位：万元）")
            self._add_kv("总资产", self._get_value("total_assets"))
            self._add_kv("总负债", self._get_value("total_liabilities"))
            self._add_kv("所有者权益", self._get_value("total_equity"))
            self._add_kv("流动资产", self._get_value("current_assets"))
            self._add_kv("流动负债", self._get_value("current_liabilities"))
            self._add_missing("【待人工粘贴：完整资产负债表】")

        # 利润表
        if is_df is not None:
            self._add_financial_table("利润表（单位：万元）", is_df)
        else:
            self._add_table_caption("利润表（单位：万元）")
            self._add_kv("营业收入", self._get_value("operating_revenue"))
            self._add_kv("营业成本", self._get_value("operating_cost"))
            self._add_kv("利润总额", self._get_value("total_profit"))
            self._add_kv("净利润", self._get_value("net_profit"))
            self._add_missing("【待人工粘贴：完整利润表】")

        # 现金流量表
        if cf_df is not None:
            self._add_financial_table("现金流量表（单位：万元）", cf_df)
        else:
            self._add_table_caption("现金流量表（单位：万元）")
            self._add_kv("经营性净现金流", self._get_value("operating_cash_flow"))
            self._add_kv("投资性净现金流", self._get_value("investing_cash_flow"))
            self._add_kv("融资性净现金流", self._get_value("financing_cash_flow"))
            self._add_missing("【待人工粘贴：完整现金流量表】")

        # (二) 收入交叉核验
        self._add_heading_2("（二）收入交叉核验")
        iv = self.data.income_verification
        if iv:
            report_rev = iv.get("report_revenue")
            bank_inflow = iv.get("bank_inflow")
            tax_revenue = iv.get("tax_revenue")
            # 用表格展示
            iv_data = {
                "数据来源": ["财务报表（利润表）", "银行流水（贷方发生额）", "纳税申报表"],
                "营业收入（万元）": [
                    report_rev or "-",
                    bank_inflow if bank_inflow else "未提供",
                    tax_revenue if tax_revenue else "未提供"
                ],
                "偏差率": [
                    "—",
                    f"{iv.get('bank_deviation_pct', 'N/A')}%" if bank_inflow else "—",
                    f"{iv.get('tax_deviation_pct', 'N/A')}%" if tax_revenue else "—"
                ],
                "核验结果": [
                    "基准值",
                    iv.get("bank_result", "未核验"),
                    iv.get("tax_result", "未核验")
                ],
            }
            self._add_financial_table("收入交叉核验", pd.DataFrame(iv_data))
            self._add_para(iv.get("conclusion", ""))

        # (三) 年度财务指标
        self._add_heading_2("（三）年度财务指标分析")
        self._add_annual_indicators()

        # (四) 杜邦分析
        self._add_heading_2("（四）杜邦分析")
        dupont = fm.get("dupont_analysis", {})
        if dupont:
            self._add_kv("净资产收益率(ROE)", f"{dupont.get('roe_pct', 'N/A')}%")
            self._add_kv("净利率", f"{dupont.get('net_margin_pct', 'N/A')}%")
            self._add_kv("资产周转率", f"{dupont.get('asset_turnover', 'N/A')}次")
            self._add_kv("权益乘数", dupont.get("equity_multiplier", "N/A"))
            self._add_para(f"杜邦分析：ROE = 净利率 × 资产周转率 × 权益乘数")
        else:
            self._add_missing("财务数据不足，无法进行杜邦分析")

        # (五) 财务分析结论
        self._add_heading_2("（五）财务分析结论")
        self._add_para(self._get_text("financial_metrics_summary"))
        self._add_para(self._get_text("financial_analysis_conclusion"))

    def _add_financial_table(self, title: str, df: pd.DataFrame) -> None:
        """从 DataFrame 渲染格式化 docx 表格"""
        if df is None or df.empty:
            self._add_missing(f"【待人工粘贴：{title}】")
            return

        self._add_table_caption(title)

        # 清理数据：替换 NaN
        df = df.fillna("")
        rows, cols = df.shape

        table = self.doc.add_table(rows=rows + 1, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # ── 表头 ──
        for j in range(cols):
            cell = table.cell(0, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(df.columns[j])[:30])
            run.font.name = FONT_BODY
            run.font.size = FONT_SIZE_TABLE
            run.font.bold = True
            self._set_east_asian_font(run, FONT_BODY)
            # 灰底
            shading = cell._element.get_or_add_tcPr()
            shd = __import__('lxml').etree.SubElement(shading, qn("w:shd"))
            shd.set(qn("w:fill"), "D9E2F3")

        # ── 数据行 ──
        for i in range(rows):
            for j in range(cols):
                cell = table.cell(i + 1, j)
                cell.text = ""
                p = cell.paragraphs[0]
                val = df.iloc[i, j]
                text = str(val)[:40] if val != "" else ""
                run = p.add_run(text)
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_TABLE
                self._set_east_asian_font(run, FONT_BODY)

        self.doc.add_paragraph()  # table 后空行

    def _try_extract_table_from_md(self, keywords: list[str]) -> pd.DataFrame | None:
        """从 MinerU 输出的 Markdown 表格中提取指定报表（懒加载缓存）"""
        md = self.data.markdown_content
        if not md:
            return None

        try:
            from phase2_analysis import _extract_tables_from_markdown
            if self._tables_cache is None:
                self._tables_cache = _extract_tables_from_markdown(md)
            tables = self._tables_cache
            # 按内容匹配（扫描第一列）
            candidates = []
            for name, rows in tables.items():
                first_col = " ".join([r[0] for r in rows if r])
                if any(kw in first_col for kw in keywords):
                    priority = 0 if "合并" in name else 1
                    candidates.append((priority, rows))
            if not candidates:
                return None
            candidates.sort(key=lambda x: x[0])
            rows = candidates[0][1]
            if len(rows) < 2:
                return None
            # 转为 DataFrame（首行为列名）
            max_cols = max(len(r) for r in rows)
            header = rows[0] + [""] * (max_cols - len(rows[0]))
            data = []
            for r in rows[1:]:
                data.append(r + [""] * (max_cols - len(r)))
            return pd.DataFrame(data, columns=header)
        except Exception:
            return None

    def _add_annual_indicators(self) -> None:
        """添加年度财务指标表（多年对比）"""
        ai = self.data.annual_indicators

        # ── 如果有年度数据，展示多年对比表 ──
        if ai and len(ai) >= 1:
            years = sorted(ai.keys())
            indicator_labels = [
                ("资产负债率", "debt_to_asset_ratio", "%"),
                ("流动比率", "current_ratio", ""),
                ("速动比率", "quick_ratio", ""),
                ("应收账款周转天数", "receivables_turnover_days", "天"),
                ("应付账款周转天数", "payables_turnover_days", "天"),
                ("存货周转天数", "inventory_turnover_days", "天"),
                ("销售利润率", "sales_profit_margin", "%"),
                ("毛利率", "gross_margin", "%"),
                ("刚性负债", "rigid_liabilities", "万元"),
                ("现金类资产", "cash_assets", "万元"),
                ("刚性负债净敞口", "rigid_liability_net", "万元"),
            ]

            cols = len(years) + 2  # 指标 + N年 + 趋势
            table = self.doc.add_table(rows=len(indicator_labels) + 1, cols=cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            # 表头：指标 | 2023年 | 2024年 | 2025年 | 趋势
            header_cells = ["指标"] + [f"{y}年" for y in years] + ["趋势"]
            for j, text in enumerate(header_cells):
                cell = table.cell(0, j)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_TABLE
                run.font.bold = True
                self._set_east_asian_font(run, FONT_BODY)
                shading = cell._element.get_or_add_tcPr()
                shd = __import__('lxml').etree.SubElement(shading, qn("w:shd"))
                shd.set(qn("w:fill"), "D9E2F3")

            for i, (label, key, unit) in enumerate(indicator_labels, 1):
                # 指标名
                cell = table.cell(i, 0)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(label)
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_TABLE
                self._set_east_asian_font(run, FONT_BODY)

                values = []
                for y in years:
                    yd = ai.get(y, {})
                    v = yd.get(key)
                    values.append(v)
                    cell = table.cell(i, years.index(y) + 1)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    if v is not None:
                        display = f"{round(v, 2)}{unit}"
                    else:
                        display = MISSING_MARKER
                    run = p.add_run(display)
                    run.font.name = FONT_BODY
                    run.font.size = FONT_SIZE_TABLE
                    if display == MISSING_MARKER:
                        run.font.color.rgb = COLOR_MISSING
                    self._set_east_asian_font(run, FONT_BODY)

                # 趋势
                trend = ""
                valid = [v for v in values if v is not None]
                if len(valid) >= 2:
                    if valid[-1] > valid[0] * 1.05:
                        trend = "↑"
                    elif valid[-1] < valid[0] * 0.95:
                        trend = "↓"
                    else:
                        trend = "→"
                cell = table.cell(i, cols - 1)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(trend)
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_TABLE
                self._set_east_asian_font(run, FONT_BODY)
        else:
            # ── 无年度数据时回退：展示当前单期值 ──
            fm = self.data.financial_metrics
            indicators = [
                ("资产负债率", fm.get("debt_to_asset_ratio") or fm.get("debt_to_asset_ratio_calc"), "%"),
                ("流动比率", fm.get("current_ratio") or fm.get("current_ratio_calc"), ""),
                ("速动比率", fm.get("quick_ratio"), ""),
                ("应收账款周转天数", fm.get("receivables_turnover_days"), "天"),
                ("应付账款周转天数", fm.get("payables_turnover_days"), "天"),
                ("存货周转天数", fm.get("inventory_turnover_days"), "天"),
                ("销售利润率", fm.get("sales_profit_margin"), "%"),
                ("毛利率", fm.get("gross_margin") or fm.get("gross_margin_calc"), "%"),
                ("刚性负债", fm.get("rigid_liabilities"), "万元"),
                ("现金类资产", fm.get("cash_assets"), "万元"),
                ("刚性负债净敞口", fm.get("rigid_liability_net"), "万元"),
            ]
            table = self.doc.add_table(rows=len(indicators) + 1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            for j, text in enumerate(["指标", "数值"]):
                cell = table.cell(0, j)
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(text)
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_TABLE
                run.font.bold = True
                self._set_east_asian_font(run, FONT_BODY)
                shading = cell._element.get_or_add_tcPr()
                shd = __import__('lxml').etree.SubElement(shading, qn("w:shd"))
                shd.set(qn("w:fill"), "D9E2F3")
            for i, (label, value, unit) in enumerate(indicators, 1):
                display = f"{value}{unit}" if value is not None else MISSING_MARKER
                for j, text in enumerate([label, display]):
                    cell = table.cell(i, j)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(str(text))
                    run.font.name = FONT_BODY
                    run.font.size = FONT_SIZE_TABLE
                    if display == MISSING_MARKER:
                        run.font.color.rgb = COLOR_MISSING
                    self._set_east_asian_font(run, FONT_BODY)

    def _add_chapter_4(self) -> None:
        """四、申请人信用状况"""
        self._add_heading_1("四、申请人信用状况")
        self._add_heading_2("（一）总体信用情况")
        self._add_para(self._get_text("overall_credit_status"))
        self._add_heading_2("（二）银行融资情况")
        self._add_para(self._get_text("bank_financing"))
        self._add_heading_2("（三）其他融资情况")
        self._add_para("非银行融资信息未采集。")

    def _add_chapter_5(self) -> None:
        """五、行业地位比较分析"""
        self._add_heading_1("五、行业地位比较分析")
        self._add_heading_2("（一）行业地位")
        self._add_para(self._get_text("industry_position"))
        self._add_heading_2("（二）竞争优势")
        self._add_para(self._get_text("competitive_advantages"))
        self._add_heading_2("（三）竞争劣势")
        self._add_para(self._get_text("competitive_disadvantages"))
        self._add_heading_2("（四）行业发展趋势")
        self._add_para(self._get_text("industry_trend"))
        self._add_heading_2("（五）价格走势分析")
        self._add_para(self._get_text("price_trend"))

    def _add_chapter_6(self) -> None:
        """六、其他重要事项"""
        self._add_heading_1("六、其他重要事项")
        self._add_heading_2("（一）诉讼及重大负面信息")
        self._add_para(self._get_text("litigation_events"))

    def _add_chapter_7(self) -> None:
        """七、授信用途及还款来源"""
        self._add_heading_1("七、授信用途及还款来源")
        self._add_heading_2("（一）授信用途")
        self._add_para(self._get_text("credit_usage"))
        self._add_para(self._get_text("credit_usage_analysis"))
        self._add_heading_2("（二）还款来源")
        self._add_para(self._get_text("repayment_source"))
        self._add_heading_2("（三）还款方式")
        self._add_para(self._get_text("repayment_method"))

    def _add_chapter_8(self) -> None:
        """八、担保情况"""
        self._add_heading_1("八、担保情况")
        gt = self.data.guarantee_types
        if not gt:
            self._add_missing("企业未提供担保资料，担保情况待补充。")
            return

        if "legal" in gt:
            self._add_heading_2("（一）法人保证担保")
            self._add_para(self._get_text("collateral_pledge"))
        if "collateral" in gt:
            self._add_heading_2("（二）抵质押担保")
            self._add_para(self._get_text("collateral_pledge"))
        if "natural_person" in gt:
            self._add_heading_2("（三）自然人保证担保")
            self._add_para(self._get_text("collateral_pledge"))

        self._add_heading_2("担保综合评价")
        self._add_para(self._get_text("guarantee_evaluation"))

    def _add_chapter_9(self) -> None:
        """九、授信收益与风险分析"""
        self._add_heading_1("九、授信收益与风险分析")
        self._add_heading_2("（一）收益分析")
        self._add_para(self._get_text("return_analysis"))
        self._add_heading_2("（二）风险评价")
        self._add_para(self._get_text("risk_evaluation"))
        self._add_heading_2("（三）风险等级")
        self._add_para(self._get_text("risk_level"), bold=True)
        self._add_heading_2("（四）风险缓释措施")
        self._add_para(self._get_text("risk_mitigation_measures"))

    def _add_chapter_10(self) -> None:
        """十、授信调查结论和授信方案"""
        self._add_heading_1("十、授信调查结论和授信方案")
        self._add_heading_2("（一）调查结论")
        self._add_para(self._get_text("investigation_conclusion"))
        self._add_heading_2("（二）上报意见")
        self._add_para(self._get_text("reported_opinion"))
        self._add_heading_2("（三）授信方案建议")
        self._add_kv("建议授信品种", self._get_text("recommended_credit_type"))
        self._add_kv("建议授信金额", self._get_text("recommended_credit_amount"))
        self._add_kv("建议授信期限", self._get_text("recommended_term"))
        self._add_kv("建议担保方式", self._get_text("recommended_guarantee"))

    # ── 主入口 ─────────────────────────────────────────────────

    def build(self) -> Document:
        """构建完整报告，返回 Document 对象"""
        self._setup_page()
        self._add_cover_page()
        self._add_chapter_1()
        self.doc.add_page_break()
        self._add_chapter_2()
        self.doc.add_page_break()
        self._add_chapter_3()
        self.doc.add_page_break()
        self._add_chapter_4()
        self.doc.add_page_break()
        self._add_chapter_5()
        self.doc.add_page_break()
        self._add_chapter_6()
        self.doc.add_page_break()
        self._add_chapter_7()
        self.doc.add_page_break()
        self._add_chapter_8()
        self.doc.add_page_break()
        self._add_chapter_9()
        self.doc.add_page_break()
        self._add_chapter_10()
        return self.doc

    def save(self, output_path: str | Path) -> None:
        """构建并保存报告"""
        doc = self.build()
        doc.save(str(output_path))
        _safe_print(f"[DocxReportBuilder] Report saved: {output_path}")
