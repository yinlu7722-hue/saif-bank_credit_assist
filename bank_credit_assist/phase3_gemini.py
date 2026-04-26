"""
phase3_gemini.py
Phase 3: Gemini 报告生成（JSON Schema + 分章节生成）
"""
from __future__ import annotations

import json
import os
import asyncio
import re
from pathlib import Path
from typing import Any, Optional

import google.generativeai as genai
from pydantic import BaseModel, Field
from docx import Document

from shared.config import GEMINI_API_KEY, GEMINI_MODEL, HTTPS_PROXY
from shared.data_schema import (
    FinalReport, ReportChapter,
    CompanyBasicInfo, BusinessAnalysis, FinancialAnalysis,
    CreditStatus, IndustryAnalysis, OtherImportantMatters,
    CreditUsageRepayment, GuaranteeInfo, RiskReturnAnalysis,
    ConclusionRecommendation,
)

# 配置代理
if HTTPS_PROXY:
    os.environ["HTTPS_PROXY"] = HTTPS_PROXY

# 配置 API Key
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)


# ============================================================================
# Prompt 模板（分章节）
# ============================================================================

CHAPTER_SYSTEM_PROMPT: str = """你是一个银行对公信贷评估报告生成助手。

【强制要求】所有财务数据、风险结论必须附带数据来源，格式如下：
{{"source_file": "文件名.扩展名", "source_location": "第X个表格/第X页", "confidence_score": 0.95}}

【禁止】
- 不要臆测未提供的信息
- 不要编造数据（如无法从输入中找到的数据，标注"未提供"）
- 不要在数据来源字段填写不确定的信息
"""

CHAPTER_PROMPTS: dict[int, tuple[str, type[BaseModel]]] = {
    1: ("基于以下企业资料，生成【一、申请人基本信息】的JSON数据。", CompanyBasicInfo),
    2: ("基于以下企业资料，生成【二、申请人经营情况】的JSON数据，包含科技型中小企业特征。", BusinessAnalysis),
    3: ("基于以下财务指标和Markdown资料，生成【三、申请人财务状况】的JSON数据。", FinancialAnalysis),
    4: ("基于以下企业信用资料，生成【四、申请人信用状况】的JSON数据。", CreditStatus),
    5: ("基于以下企业资料，生成【五、行业地位比较分析】的JSON数据。", IndustryAnalysis),
    6: ("基于以下企业资料，生成【六、其他重要事项】的JSON数据。", OtherImportantMatters),
    7: ("基于以下授信申请资料，生成【七、授信用途及还款来源】的JSON数据。", CreditUsageRepayment),
    8: ("基于以下担保资料，生成【八、担保情况】的JSON数据。", GuaranteeInfo),
    9: ("基于以下财务和合规分析结果，生成【九、授信收益与风险分析】的JSON数据。", RiskReturnAnalysis),
    10: ("基于以上所有分析结果，生成【十、授信调查结论和授信方案】的JSON数据。", ConclusionRecommendation),
}


# ============================================================================
# Gemini 客户端封装
# ============================================================================

class GeminiReportGenerator:
    """
    Gemini 报告生成器（JSON Schema 结构化输出）
    使用 Pydantic 模型 + response_schema 强制结构化输出，
    十大章节分步生成，最终映射到 docx 模板。
    """

    def __init__(self, api_key: str | None = None) -> None:
        key: str = api_key or GEMINI_API_KEY
        if key:
            genai.configure(api_key=key)
        self.model = genai.GenerativeModel(GEMINI_MODEL)

    async def generate_chapter(
        self,
        chapter_number: int,
        enterprise_data: dict,
        financial_metrics: dict,
        compliance_results: dict,
    ) -> ReportChapter:
        """
        生成单个报告章节（JSON Schema 强制结构化输出）
        """
        system_prompt = CHAPTER_SYSTEM_PROMPT
        user_prompt_template, pydantic_model = CHAPTER_PROMPTS[chapter_number]

        user_prompt: str = f"""
{user_prompt_template}

## 输入数据

### 企业资料
{json.dumps(enterprise_data, ensure_ascii=False, indent=2)}

### 财务指标
{json.dumps(financial_metrics, ensure_ascii=False, indent=2)}

### 合规筛查结果
{json.dumps(compliance_results, ensure_ascii=False, indent=2)}

## 输出要求
严格按以下JSON Schema格式输出（必须包含所有字段，不要省略）:

{json.dumps(pydantic_model.model_json_schema(), ensure_ascii=False, indent=2)}

【重要】只输出JSON，不要有其他文字。
"""

        try:
            response = await asyncio.to_thread(
                self.model.generate_content,
                contents=user_prompt,
                generation_config={
                    "system_instruction": system_prompt,
                    "response_mime_type": "application/json",
                    "response_schema": pydantic_model.model_json_schema(),
                },
            )

            raw_text: str = response.text
            chapter_data: dict = json.loads(raw_text)

            return ReportChapter(
                chapter_number=chapter_number,
                chapter_name=self._get_chapter_name(chapter_number),
                content=chapter_data,
                generation_status="success",
            )

        except Exception as e:
            return ReportChapter(
                chapter_number=chapter_number,
                chapter_name=self._get_chapter_name(chapter_number),
                content={},
                generation_status="failed",
                error_message=str(e),
            )

    async def generate_full_report(
        self,
        enterprise_data: dict,
        financial_metrics: dict,
        compliance_results: dict,
    ) -> FinalReport:
        """
        生成完整报告（十大章节分步串行生成）
        """
        print("\n" + "=" * 60)
        print("开始生成报告（分章节进行，共10章）...")
        print("=" * 60 + "\n")

        chapters: list[ReportChapter] = []
        failed_count: int = 0

        # 十大章节串行生成（避免token溢出）
        for chapter_num in range(1, 11):
            print(f"  [Chapter {chapter_num}/10] 生成中: {self._get_chapter_name(chapter_num)}...")

            chapter: ReportChapter = await self.generate_chapter(
                chapter_num,
                enterprise_data,
                financial_metrics,
                compliance_results,
            )
            chapters.append(chapter)

            if chapter.generation_status == "success":
                print(f"    ✓ Chapter {chapter_num} 生成成功")
            else:
                failed_count += 1
                print(f"    ✗ Chapter {chapter_num} 生成失败: {chapter.error_message}")

            # Gemini 限速保护：每章间隔1秒
            await asyncio.sleep(1.0)

        # 综合风险等级判定
        overall_risk: str = self._determine_risk_level(chapters)

        # 提取需人工核验项
        manual_verify: list[str] = [
            ch.chapter_name
            for ch in chapters
            if ch.content.get("data_source", {}).get("source_file") in ("未知", "未提供")
            or ch.generation_status == "failed"
        ]

        report: FinalReport = FinalReport(
            enterprise_name=enterprise_data.get("enterprise_name", "未知"),
            report_date=enterprise_data.get("report_date", ""),
            overall_risk_level=overall_risk,
            chapters=chapters,
            compliance_status=compliance_results.get("overall", "UNKNOWN"),
            manual_verify_items=manual_verify,
        )

        print("\n" + "=" * 60)
        print(f"报告生成完成！成功: {10 - failed_count}/10, 失败: {failed_count}/10")
        print(f"综合风险等级: {overall_risk}")
        print("=" * 60 + "\n")

        return report

    def _get_chapter_name(self, number: int) -> str:
        names: dict[int, str] = {
            1: "一、申请人基本信息",
            2: "二、申请人经营情况",
            3: "三、申请人财务状况",
            4: "四、申请人信用状况",
            5: "五、行业地位比较分析",
            6: "六、其他重要事项",
            7: "七、授信用途及还款来源",
            8: "八、担保情况",
            9: "九、授信收益与风险分析",
            10: "十、授信调查结论和授信方案",
        }
        return names.get(number, f"第{number}章")

    def _determine_risk_level(self, chapters: list[ReportChapter]) -> str:
        """根据各章节内容综合判定风险等级"""
        if any(ch.generation_status == "failed" for ch in chapters):
            return "高风险（部分章节生成失败）"
        return "待评估"


# ============================================================================
# 模板填充引擎（防呆版）
# ============================================================================

class ReportTemplateFiller:
    """
    Word 模板渲染器（防呆版）
    【核心原则】
    1. 绝不直接对 paragraph.text 赋值（会丢失所有原有样式）
    2. 必须在 Run 层级进行替换
    3. 占位符（如 {{ROE}}）应在模板中连续输入，确保在同一个 Run 内
    4. 支持普通段落 + 表格 + 嵌套表格的深度遍历
    """

    PLACEHOLDER_PATTERN: re.Pattern = re.compile(r"\{\{([^}]+)\}\}")

    # 过滤溯源相关字段（Word 报告必须纯净）
    AUDIT_FIELDS: set[str] = {
        "source_file", "source_location", "confidence_score",
        "data_source", "source", "confidence", "provenance"
    }

    def __init__(self, template_path: str | Path) -> None:
        self.template_path: Path = Path(template_path)

    def _replace_text_in_runs(
        self,
        paragraph,
        data_dict: dict[str, Any],
    ) -> None:
        """安全替换段落中的占位符，最大程度保留 Word 原本的字体、字号和加粗等样式。"""
        for key, value in data_dict.items():
            placeholder: str = f"{{{{{key}}}}}"

            if placeholder not in paragraph.text:
                continue

            # 格式化值
            if value is None:
                display_value: str = ""
            elif isinstance(value, (int, float)):
                display_value = str(value)
            elif isinstance(value, dict):
                display_value = json.dumps(value, ensure_ascii=False, indent=1)
            elif isinstance(value, list):
                display_value = json.dumps(value, ensure_ascii=False)
            else:
                display_value = str(value)

            # 策略1：逐个 Run 替换
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, display_value)

            # 策略2（防呆终极手段）：若占位符被拆分到不同 Run
            if placeholder in paragraph.text:
                full_text: str = paragraph.text
                new_text: str = full_text.replace(placeholder, display_value)
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for i in range(1, len(paragraph.runs)):
                        paragraph.runs[i].text = ""
                else:
                    paragraph.add_run(new_text)

    def _process_paragraphs(self, paragraphs, data_dict: dict[str, Any]) -> None:
        """遍历并替换一组段落中的占位符"""
        for paragraph in paragraphs:
            if paragraph.text.strip():
                self._replace_text_in_runs(paragraph, data_dict)

    def _process_table(self, table, data_dict: dict[str, Any]) -> None:
        """递归遍历表格（含嵌套表格）并替换占位符"""
        for row in table.rows:
            for cell in row.cells:
                self._process_paragraphs(cell.paragraphs, data_dict)
                if hasattr(cell, "tables") and cell.tables:
                    for nested_table in cell.tables:
                        self._process_table(nested_table, data_dict)

    def fill(self, report: FinalReport, output_path: str | Path) -> None:
        """将报告 JSON 填充到 Word 模板（防呆版）"""
        try:
            doc: Document = Document(str(self.template_path))
        except Exception as e:
            raise ValueError(f"模板加载失败，请确保模板是 .docx 格式。错误信息: {str(e)}")

        output: Path = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)

        # 将 FinalReport 结构展平为 {key: value} 字典
        data_dict: dict[str, Any] = self._flatten_report(report)

        # 1. 替换普通段落中的占位符
        self._process_paragraphs(doc.paragraphs, data_dict)

        # 2. 深度遍历替换所有表格中的占位符
        for table in doc.tables:
            self._process_table(table, data_dict)

        # 保存
        doc.save(str(output))
        print(f"报告已成功生成并保留原有格式: {output}")

    def _flatten_report(self, report: FinalReport) -> dict[str, Any]:
        """将 FinalReport 展平为单层 key-value 字典，用于占位符替换"""
        data: dict[str, Any] = {}

        # 基本信息
        data["enterprise_name"] = report.enterprise_name
        data["report_date"] = report.report_date
        data["overall_risk_level"] = report.overall_risk_level
        data["compliance_status"] = report.compliance_status

        # 各章节内容
        for chapter in report.chapters:
            if chapter.generation_status == "success":
                flat_content: dict[str, Any] = self._flatten_dict(chapter.content)
                data.update(flat_content)

        # 过滤所有溯源相关字段（Word 报告必须纯净）
        return {
            k: v for k, v in data.items()
            if not any(aud in k.lower() for aud in self.AUDIT_FIELDS)
        }

    def _flatten_dict(self, d: dict | Any, parent_key: str = "") -> dict[str, Any]:
        """将嵌套字典展平为单层 key-value"""
        items: list[tuple[str, Any]] = []
        for k, v in d.items():
            new_key: str = f"{parent_key}.{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key).items())
            elif isinstance(v, list):
                items.append((new_key, json.dumps(v, ensure_ascii=False)))
            else:
                items.append((new_key, v))
        return dict(items)


# ============================================================================
# 便捷封装
# ============================================================================

async def generate_report(
    enterprise_data: dict,
    financial_metrics: dict,
    compliance_results: dict,
    template_path: str | Path,
    output_path: str | Path,
) -> FinalReport:
    """
    端到端报告生成（防呆版）
    """
    # 初始化生成器
    generator: GeminiReportGenerator = GeminiReportGenerator()

    # 分章节生成
    report: FinalReport = await generator.generate_full_report(
        enterprise_data,
        financial_metrics,
        compliance_results,
    )

    # 填充模板（防呆版：Run 层级替换，保留原有样式）
    filler: ReportTemplateFiller = ReportTemplateFiller(template_path)
    filler.fill(report, output_path)

    return report
