"""
phase3_report.py
Phase 3: 报告生成 — DocxReportBuilder 从零构建 .docx
（保留 TableEngine 用于 YAML 驱动的动态表格，已弃用模板填充）
"""
from __future__ import annotations

import re
from datetime import datetime, timezone, timedelta
from pathlib import Path
from typing import Any

import pandas as pd
import yaml
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from shared.config import OUTPUT_DIR

# ── 日志文件（UTF-8，完全避开终端 GBK 编码）──────────────────────
LOG_PATH: Path = OUTPUT_DIR / "phase3.log"
_tz = timezone(timedelta(hours=8))  # Asia/Shanghai


MAX_LOG_BYTES: int = 500 * 1024  # 500KB


def _log(msg: str) -> None:
    """UTF-8 文件日志 — 不写终端，根除 GBK 编码错误；超过阈值自动轮转"""
    try:
        if LOG_PATH.exists() and LOG_PATH.stat().st_size > MAX_LOG_BYTES:
            bak = LOG_PATH.with_suffix(".log.bak")
            bak.write_text(LOG_PATH.read_text(encoding="utf-8")[-100000:], encoding="utf-8")
            LOG_PATH.write_text("", encoding="utf-8")
        ts = datetime.now(_tz).strftime("%H:%M:%S")
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            f.write(f"[{ts}] {msg}\n")
    except Exception:
        pass  # 日志写入失败不影响主流程


def _safe_str(obj: object) -> str:
    """将任意对象转为 ASCII-safe 字符串，过滤不可编码字符"""
    return str(obj).encode("ascii", errors="replace").decode("ascii")


# ============================================================================
# 表格动态生成引擎（YAML 配置驱动 + Pandas）
# ============================================================================

class TableEngine:
    """
    表格引擎：读取 YAML 配置 → 从数据字典取数 → pandas DataFrame → docx 表格
    支持 5 种数据源：phase2_flat / phase2_list / phase1_markdown / phase25_text / static_fallback
    """

    def __init__(self, yaml_path: str | Path | None = None) -> None:
        if yaml_path is None:
            yaml_path = Path(__file__).parent / "tables_config.yaml"
        with open(yaml_path, "r", encoding="utf-8") as f:
            self.config = yaml.safe_load(f)
        self.tables = self.config.get("tables", {})

    # ── 主入口：填充所有表格 ─────────────────────────────────

    def fill_all_tables(self, doc: Document, data_dict: dict[str, Any]) -> int:
        """
        扫描文档中所有 {{TABLE:table_name}} 标记，替换为动态生成的表格。
        返回成功填充的表格数。
        """
        filled: int = 0

        for para in list(doc.paragraphs):
            match = re.match(r"\{\{TABLE:(\w+)\}\}", para.text.strip())
            if not match:
                continue

            table_name = match.group(1)
            table_config = self.tables.get(table_name)
            if not table_config:
                _log(f"[TableEngine] Unknown table: {table_name}, skipping")
                continue

            try:
                df = self._extract(table_name, table_config, data_dict)
                self._insert_table_at_paragraph(doc, para, df, table_config)
                # 删除标记段落
                p_element = para._element
                p_element.getparent().remove(p_element)
                filled += 1
                _log(f"[TableEngine] {table_name}: {len(df)} rows generated")
            except Exception as e:
                _log(f"[TableEngine] {table_name} ERROR: {_safe_str(e)}")
                # 失败时替换为 fallback 文字
                fallback = table_config.get("fallback", f"【待人工补充：{table_name}】")
                para.text = fallback

        return filled

    # ── 数据提取 ────────────────────────────────────────────

    def _extract(self, name: str, cfg: dict, data: dict) -> pd.DataFrame:
        """根据 source 类型从数据字典中提取 DataFrame"""
        source = cfg.get("source", "static_fallback")

        if source == "phase2_flat":
            return self._extract_flat(name, cfg, data)
        elif source == "phase2_list":
            return self._extract_list(name, cfg, data)
        elif source == "phase1_markdown":
            return self._extract_markdown(name, cfg, data)
        elif source == "static_fallback":
            return self._extract_fallback(cfg)
        else:
            return self._extract_fallback(cfg)

    def _extract_flat(self, name: str, cfg: dict, data: dict) -> pd.DataFrame:
        """dict → 键值对 DataFrame（phase2_flat）"""
        source_key = cfg.get("source_key", "")
        source_data = data.get(source_key, {}) or {}

        row_labels = cfg.get("row_labels", {})
        if row_labels:
            rows = []
            for key, label in row_labels.items():
                value = source_data.get(key, "")
                if value is None:
                    value = "【待提取】"
                elif isinstance(value, (int, float)):
                    fmt = cfg.get("columns", [{}])[1].get("format", "")
                    if fmt and isinstance(value, float):
                        value = f"{value:{fmt}}"
                    else:
                        value = str(value)
                else:
                    value = str(value)
                rows.append({"field": label, "value": value})

            # 额外行（如经营资质中的"其他资质"）
            extra = cfg.get("extra_rows", [])
            for extra_label in extra:
                rows.append({"field": extra_label, "value": "【待人工补充】"})

            return pd.DataFrame(rows)

        # 无 row_labels 时，将整个 dict 转为键值对
        items = [[str(k), str(v)] for k, v in source_data.items() if v is not None]
        if not items:
            items = [["【待人工补充】", ""]]
        return pd.DataFrame(items, columns=["field", "value"])

    def _extract_list(self, name: str, cfg: dict, data: dict) -> pd.DataFrame:
        """list[dict] → DataFrame（phase2_list）"""
        source_key = cfg.get("source_key", "")
        source_data = data.get(source_key, [])

        # 支持 JSON 字符串
        if isinstance(source_data, str):
            try:
                source_data = json.loads(source_data)
            except json.JSONDecodeError:
                source_data = []

        if not source_data or not isinstance(source_data, list):
            return self._extract_fallback(cfg)

        df = pd.DataFrame(source_data)

        # 列名中文化
        columns_cfg = cfg.get("columns", [])
        rename_map = {c["key"]: c["label"] for c in columns_cfg if "key" in c}
        df = df.rename(columns=rename_map)

        # 只保留配置中的列
        keep_cols = [c["label"] for c in columns_cfg if c.get("label") in df.columns]
        if keep_cols:
            df = df[keep_cols]

        df = df.fillna("")
        return df

    def _extract_markdown(self, name: str, cfg: dict, data: dict) -> pd.DataFrame:
        """Phase1 Markdown → pd.read_html() → DataFrame（phase1_markdown）"""
        markdown = data.get("markdown_content", "") or ""
        if not markdown:
            return self._extract_fallback(cfg)

        try:
            tables = pd.read_html(markdown)
            idx = cfg.get("table_index", 0)
            if idx < len(tables):
                df = tables[idx]
                # 清理列名
                df.columns = [str(c).strip() for c in df.columns]
                return df.fillna("")
        except Exception:
            pass

        return self._extract_fallback(cfg)

    def _extract_fallback(self, cfg: dict) -> pd.DataFrame:
        """无数据源 — 返回含 fallback 提示的空 DataFrame"""
        columns = cfg.get("columns", [])
        if columns:
            col_labels = [c.get("label", c.get("key", "")) for c in columns]
            fallback_msg = cfg.get("fallback", "【待人工补充】")
            return pd.DataFrame([[fallback_msg] + [""] * (len(col_labels) - 1)], columns=col_labels)
        return pd.DataFrame([["【待人工补充】"]], columns=["备注"])

    # ── 格式化 ──────────────────────────────────────────────

    def _format_dataframe(self, df: pd.DataFrame, cfg: dict) -> pd.DataFrame:
        """对 DataFrame 应用列格式化规则"""
        columns_cfg = cfg.get("columns", [])
        for col_cfg in columns_cfg:
            label = col_cfg.get("label", "")
            fmt = col_cfg.get("format", "")
            if fmt and label in df.columns:
                try:
                    df[label] = pd.to_numeric(df[label], errors="coerce")
                    df[label] = df[label].apply(
                        lambda x: f"{x:{fmt}}" if pd.notna(x) else ""
                    )
                except Exception:
                    pass
        return df

    # ── 插入 docx 表格 ──────────────────────────────────────

    def _insert_table_at_paragraph(
        self, doc: Document, para, df: pd.DataFrame, cfg: dict
    ) -> None:
        """在指定段落位置插入 pandas DataFrame 渲染的 docx 表格"""
        df = self._format_dataframe(df, cfg)

        rows = len(df) + 1  # +1 for header
        cols = len(df.columns)
        table = doc.add_table(rows=rows, cols=cols, style="Table Grid")

        # 写表头
        for j, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col_name)
            self._set_cell_style(cell, bold=True, size=9, bg_color="D9E2F3")

        # 写数据行
        for i, (_, row) in enumerate(df.iterrows()):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(val) if pd.notna(val) else ""
                self._set_cell_style(cell, bold=False, size=9)

        # 设置列宽
        columns_cfg = cfg.get("columns", [])
        for j, col_cfg in enumerate(columns_cfg):
            if j < cols:
                width_pct = col_cfg.get("width", 10)
                # 按比例换算（总宽度约16cm）
                width_cm = width_pct * 0.16
                for row in table.rows:
                    row.cells[j].width = Cm(width_cm)

        # 在表格前插入题注
        caption = cfg.get("caption", "")
        if caption:
            caption_para = para.insert_paragraph_before(caption)
            caption_run = caption_para.runs[0] if caption_para.runs else caption_para.add_run(caption)
            caption_run.bold = True
            caption_run.font.size = Pt(10)

        # 表格后加空行
        para.insert_paragraph_before("")

    @staticmethod
    def _set_cell_style(cell, bold: bool = False, size: int = 9, bg_color: str | None = None) -> None:
        """设置单元格字体样式"""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(size)
                run.font.name = "Microsoft YaHei"
                run.bold = bold
            if not paragraph.runs and cell.text:
                # 确保新文本也应用样式
                pass
        if bg_color:
            shading_elm = cell._element.get_or_add_tcPr()
            shading = shading_elm.makeelement(qn("w:shd"), {
                qn("w:fill"): bg_color,
                qn("w:val"): "clear",
            })
            shading_elm.append(shading)


# ============================================================================
# 便捷封装
# ============================================================================

async def generate_report(
    enterprise_name: str,
    basic_info: dict,
    financial_metrics: dict,
    tech_metrics: dict,
    inference_text: dict,
    income_verification: dict,
    guarantee_types: list[str],
    markdown_content: str,
    output_path: str | Path,
    annual_indicators: dict | None = None,
) -> None:
    """端到端报告生成：数据组装 → DocxReportBuilder → 保存"""
    from docx_builder import ReportData, DocxReportBuilder

    data = ReportData(
        enterprise_name=enterprise_name,
        basic_info=basic_info or {},
        financial_metrics=financial_metrics or {},
        tech_metrics=tech_metrics or {},
        inference_text=inference_text or {},
        income_verification=income_verification or {},
        guarantee_types=guarantee_types or [],
        markdown_content=markdown_content or "",
        annual_indicators=annual_indicators or {},
    )

    builder = DocxReportBuilder(data)
    builder.save(output_path)
    _log(f"Report saved: {output_path}")
