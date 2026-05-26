#!/usr/bin/env python3
# -*- coding: utf-8 -*-
'''生成技术说明文档（Word格式）— 银行对公信贷智能化辅助系统'''
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

# ── 样式常量 ──
FONT_TITLE = '黑体'
FONT_H1 = '黑体'
FONT_H2 = '黑体'
FONT_H3 = '黑体'
FONT_BODY = '宋体'

TITLE_SIZE = Pt(22)
H1_SIZE = Pt(16)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)
BODY_SIZE = Pt(12)

COLOR_TITLE = RGBColor(0x00, 0x33, 0x66)
COLOR_H1 = RGBColor(0x00, 0x33, 0x66)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY = RGBColor(0x33, 0x33, 0x33)

# ── 辅助函数 ──
def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_para(doc, text, font_name=FONT_BODY, font_size=BODY_SIZE, bold=False,
             color=COLOR_BLACK, alignment=None, space_after=Pt(6),
             first_line_indent=None, line_spacing=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    run.font.color.rgb = color
    return p

def add_h1(doc, text):
    return add_para(doc, text, font_name=FONT_H1, font_size=H1_SIZE,
                    bold=True, color=COLOR_H1, space_after=Pt(12), line_spacing=1.5)

def add_h2(doc, text):
    return add_para(doc, text, font_name=FONT_H2, font_size=H2_SIZE,
                    bold=True, color=COLOR_H1, space_after=Pt(8), line_spacing=1.5)

def add_h3(doc, text):
    return add_para(doc, text, font_name=FONT_H3, font_size=H3_SIZE,
                    bold=True, color=COLOR_BLACK, space_after=Pt(6), line_spacing=1.5)

def add_body(doc, text, indent=True):
    return add_para(doc, text, font_name=FONT_BODY, font_size=BODY_SIZE,
                    first_line_indent=Cm(0.74) if indent else None,
                    line_spacing=1.5, space_after=Pt(6))

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_GRAY
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.74)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        run.font.size = BODY_SIZE
        run.bold = True
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    run.font.size = BODY_SIZE
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '003366')

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            run.font.size = Pt(10)
            if c == 0:
                run.bold = True
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F6FC')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table

# ========================================================================
# 封面
# ========================================================================
for _ in range(6):
    doc.add_paragraph()

add_para(doc, '银行对公信贷流程智能化辅助系统', font_name=FONT_TITLE,
         font_size=TITLE_SIZE, bold=True, color=COLOR_TITLE,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

add_para(doc, '技术说明文档', font_name=FONT_H1, font_size=H1_SIZE,
         bold=True, color=COLOR_TITLE,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))

add_para(doc, '版本 v5.5', font_name=FONT_BODY, font_size=Pt(12),
         color=COLOR_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
add_para(doc, '日期: ' + datetime.date.today().strftime('%Y年%m月%d日'),
         font_name=FONT_BODY, font_size=Pt(12),
         color=COLOR_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
add_para(doc, 'AI4Leader 课程实践结课作业',
         font_name=FONT_BODY, font_size=Pt(12),
         color=COLOR_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))

doc.add_page_break()

# ========================================================================
# 目录页
# ========================================================================
add_h1(doc, '目  录')
doc.add_paragraph()

toc_items = [
    ('一、项目背景', False),
    ('    1.1 业务背景与痛点', True),
    ('    1.2 项目目标', True),
    ('    1.3 系统边界与定位', True),
    ('二、系统架构', False),
    ('    2.1 整体架构设计', True),
    ('    2.2 四阶段流水线', True),
    ('    2.3 技术栈总览', True),
    ('    2.4 数据流与状态机', True),
    ('    2.5 模块依赖关系', True),
    ('三、核心算法', False),
    ('    3.1 财务报表智能提取算法', True),
    ('    3.2 收入交叉验证算法', True),
    ('    3.3 杜邦分析与财务指标体系', True),
    ('    3.4 多智能体辩论系统（贷审会模拟）', True),
    ('    3.5 AI推理并行批量调度', True),
    ('四、关键技术选型与实现难点', False),
    ('    4.1 技术选型说明', True),
    ('    4.2 实现难点与解决方案', True),
    ('五、总结与展望', False),
]

for item, is_sub in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.8
    run = p.add_run(item)
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    run.font.size = Pt(11)
    if not is_sub:
        run.bold = True

doc.add_page_break()

# ========================================================================
# 一、项目背景
# ========================================================================
add_h1(doc, '一、项目背景')

add_h2(doc, '1.1 业务背景与痛点')

add_body(doc,
    '银行对公信贷业务是商业银行的核心资产业务，其审批质量直接影响银行的资产质量和经营效益。'
    '一笔典型的中型企业授信业务涉及企业提交的各类资料多达20余项，包括营业执照、公司章程、近三年财务报表、'
    '银行流水、纳税申报表、专利证书、上下游合同等，涵盖PDF、图片、Excel、Word等多种格式。')

add_body(doc, '传统信贷审批流程面临以下核心痛点：')

pain_points = [
    ('信息碎片化',
     '企业资料分散在多种格式、多个文件中，信贷员需要手工翻阅、摘录关键数据，一份完整的调查报告通常需要2-3个工作日才能完成。'),
    ('财务分析工作量大',
     '财务报表科目多达数十项，需要手动计算资产负债率、流动比率、速动比率、毛利率、净利率等十余项关键指标，并进行杜邦分析和多年趋势对比，繁琐且易出错。'),
    ('经验依赖度高',
     '对企业信用风险的综合判断高度依赖信贷员的个人经验，不同信贷员对同一企业的评价可能存在显著差异，缺少标准化的评估框架。'),
    ('审批流程长',
     '从贷前调查、报告撰写、风险评审到贷审会表决，全流程涉及多个角色和多轮沟通，信息传递效率低。'),
    ('监管合规压力',
     '监管要求贷前调查须做到“双人实地、尽职调查”，调查报告需覆盖申请人基本信息、经营情况、财务状况、信用状况、行业分析、担保评价、风险缓释等十大板块，内容要求全面且规范。'),
]

for title, desc in pain_points:
    add_bullet(doc, desc, bold_prefix='● ' + title + '：')

add_h2(doc, '1.2 项目目标')

add_body(doc,
    '本项目旨在利用大语言模型（LLM）、多模态文档解析、多智能体协作等AI技术，构建一套覆盖对公信贷全流程的'
    '智能化辅助系统，实现以下目标：')

goals = [
    '自动化资料解析：将企业提交的PDF、Word、Excel、图片等多格式资料自动转换为结构化文本，准确提取营业执照信息、财务报表科目、银行流水数据等关键字段。',
    '智能化财务分析：基于提取的结构化数据，自动完成11项财务比率计算、杜邦分析、多年趋势对比和收入交叉验证，输出专业的财务分析结论。',
    'AI辅助报告生成：利用大语言模型对42个定性分析字段（如经营模式、行业地位、风险评估、授信建议等）进行专业推理和文本生成，最终从零构建一份符合银行规范的Word格式授信调查报告。',
    '多智能体贷审会模拟：构建5位不同角色的AI评审官，通过5轮结构化辩论模拟真实贷审会流程，输出终审结论和授信建议，为信贷决策提供多维度参考。',
]

for g in goals:
    add_bullet(doc, g)

add_h2(doc, '1.3 系统边界与定位')

add_body(doc,
    '本系统定位为“AI辅助决策工具”，而非“自动化审批系统”。核心设计原则包括：')

principles = [
    '人在环中（Human-in-the-Loop）：Phase2数据提取完成后，系统进入“核对”状态，信贷员可在前端编辑修正AI提取的财务数据，确认无误后再触发报告生成。AI负责效率提升，人工负责最终把关。',
    '可解释性优先：报告中的所有AI生成内容均标注数据来源（“根据所提供的财务报表推算”），数据不足时明确标注“根据现有资料无法确定，建议实地调研补充”，避免AI幻觉产生误导。',
    '模块化流水线：系统按“资料解析→数据提取→AI推理→报告生成→贷审会模拟”四阶段流水线设计，各阶段独立运行、松耦合，可单独调试和替换。',
]

for pr in principles:
    add_bullet(doc, pr)

doc.add_page_break()

# ========================================================================
# 二、系统架构
# ========================================================================
add_h1(doc, '二、系统架构')

add_h2(doc, '2.1 整体架构设计')

add_body(doc,
    '系统采用前后端分离的B/S架构。前端为纯HTML/CSS/JavaScript单页应用，后端为FastAPI异步Web服务，'
    '业务逻辑按四阶段流水线组织。整体架构如下图所示（文字描述）：')

add_body(doc,
    '前端（index.html）通过HTTP REST API与后端通信，使用UUID会话标识维护多用户并发状态。'
    '后端服务（server.py）作为编排中心，协调Phase1-Phase4各模块的执行，'
    '长耗时任务（文档解析、AI推理）通过asyncio.create_task放入后台执行，前端通过轮询API获取实时进度。')

add_h2(doc, '2.2 四阶段流水线')

add_body(doc, '系统的核心业务流程分为四个阶段，数据按阶段逐步深化：')

add_h3(doc, 'Phase 1 — 多模态文档解析')

add_body(doc,
    '输入为企业上传的各类资料文件（PDF/Word/图片/Excel/HTML），输出为统一的Markdown结构化文本和Excel表格数据。')
add_body(doc,
    '技术路线：对于PDF、Word、图片等非结构化文件，通过MinerU云API（基于VLM视觉语言模型）进行OCR识别和版面分析，'
    '输出Markdown格式的结构化文本。对于Excel文件（.xls/.xlsx/.xlsm），通过pandas本地解析为Markdown表格，'
    '确保财务报表类数据的精确提取，避免OCR引入的识别误差。MinerU客户端采用异步状态机设计，支持批量上传、'
    '指数退避重试和递增间隔轮询。')

add_h3(doc, 'Phase 2 — 数据提取与财务分析')

add_body(doc, '此阶段为核心数据处理环节，包含五个子步骤：')
add_body(doc,
    'Step 2.1 财务报表提取：从Markdown文本中定位资产负债表、利润表、现金流量表，按内容匹配报表类型（扫描首列科目名），'
    '按年份定位目标列（取最近完整财年），合并报表优先于单体报表。提取总资产、总负债、营业收入、净利润等15+个关键科目。')
add_body(doc,
    'Step 2.2 科技属性提取：提取企业的高新技术企业证书有效期、研发费用占比、专利和知识产权数量、核心技术团队规模等信息。')
add_body(doc,
    'Step 2.3 企业基本信息提取：从营业执照提取统一社会信用代码、法定代表人、注册资本、注册日期；'
    '从公司章程提取股东结构和实际控制人；识别担保类型（法人担保/抵质押担保/自然人担保）。')
add_body(doc,
    'Step 2.4 财务指标计算（纯Python计算，无AI调用）：基于提取的原始科目数据，计算11项年度财务指标'
    '（资产负债率、流动比率、速动比率、应收账款周转天数等），并进行杜邦分析（ROE=净利率×资产周转率×权益乘数）。')
add_body(doc,
    'Step 2.5 收入交叉验证：将财务报表营业收入与银行流水（贷方发生额）、纳税申报表（申报营业收入）三方交叉比对，'
    '按三级阈值判定数据一致性（吻合/偏差/重大偏差）。')

add_h3(doc, 'Phase 3 — AI推理与报告生成')

add_body(doc,
    'Phase 3.1 AI智能推理（phase2_inference.py）：将42个需要专业判断的定性分析字段（如主营业务描述、'
    '管理层评价、经营风险分析、行业地位评估、授信建议等）分为9个批次，通过asyncio.Semaphore(3)控制并发，'
    '并行调用Minimax API（Anthropic兼容接口，模型为claude-3.5-sonnet-20240620），总计12次API调用完成全部推理。')
add_body(doc,
    'Phase 3.2 报告生成（docx_builder.py）：完全舍弃Word模板方案，使用python-docx从零程序化构建文档。'
    'DocxReportBuilder类负责10个章节的排版与内容填充，支持A4页面、黑体/宋体混排、表格渲染、'
    '缺失数据红色标注等功能。同时通过TableEngine（table_config.yaml驱动）实现可配置的财务表格动态渲染。')

add_h3(doc, 'Phase 4 — 多智能体贷审会模拟（可选）')

add_body(doc,
    '构建5位AI评委（牵头审批官、风险审批官、行业审批官、财务审批官、合规审批官），每位评委拥有独立的'
    '角色设定、评审标准和性格特征。通过5轮结构化辩论（独立评审→风险合规交锋→行业财务交锋→自由辩论→终审汇总），'
    '由牵头审批官输出结构化的终审结论JSON，包含授信结论、风险提示和授信方案建议。发言顺序按确定性优先级轮换'
    '（风险→行业→财务→合规→主持），每轮发言基于前序全部发言的辩论上下文。')

add_h2(doc, '2.3 技术栈总览')

tech_headers = ['层次', '技术选型', '说明']
tech_rows = [
    ['后端框架', 'FastAPI 0.109 + Uvicorn 0.27', '原生异步支持，自动OpenAPI文档，适合AI服务的长耗时请求场景'],
    ['前端', '纯HTML5 + CSS3 + Vanilla JS', '无框架依赖，响应式设计，通过localStorage维护会话'],
    ['文档解析', 'MinerU云API（VLM模型）', '多模态OCR与版面分析，支持PDF/Word/PPT/图片，输出Markdown'],
    ['表格解析', 'pandas 2.2 + openpyxl', 'Excel文件本地解析，确保财务报表数字精确性'],
    ['LLM服务', 'Minimax Anthropic兼容API', '模型: claude-3.5-sonnet-20240620，国内可直接访问'],
    ['LLM SDK', 'anthropic 0.25（Python SDK）', '通过base_url重定向至Minimax端点实现兼容调用'],
    ['报告生成', 'python-docx 1.1', '完全程序化构建Word文档，无模板依赖，支持复杂表格、中文字体'],
    ['数据模型', 'Pydantic 2.6', '请求/响应验证，统一文档清单和校验规则定义'],
    ['异步HTTP', 'aiohttp 3.9', 'MinerU API的异步批量上传和轮询'],
    ['配置管理', 'python-dotenv + PyYAML', '.env环境变量 + YAML表格配置'],
    ['并发控制', 'asyncio.Semaphore', 'AI推理批次的并发度控制（最多3并行）'],
    ['HTML解析', 'BeautifulSoup4 + lxml', 'Markdown中HTML表格的解析回退方案'],
]

add_table(doc, tech_headers, tech_rows, col_widths=[2.5, 5.0, 7.5])

add_h2(doc, '2.4 数据流与状态机')

add_body(doc,
    '系统基于会话状态机驱动前端交互。每个浏览器会话分配一个UUID，服务端维护SessionState对象，'
    '状态按如下顺序流转：')

add_body(doc,
    'idle（等待上传）→ parsing（文档解析中）→ parsing_done（解析完成）→ '
    'analyzing（数据分析中）→ verifying（人工核对）→ generating（报告生成中）→ done（可下载）→ '
    'reviewing（可选，贷审会模拟中）。')

add_body(doc,
    '前端通过轮询 GET /api/status/{session_id} 获取当前状态和进度百分比，根据状态显示不同的操作面板。'
    '长耗时任务（MinerU解析、AI推理）使用asyncio.create_task放入后台协程执行，不阻塞HTTP响应。'
    '会话1小时无活动后自动清理。')

add_h2(doc, '2.5 模块依赖关系')

add_body(doc, '核心模块按职责分层，以下为各模块之间的导入依赖关系：')

add_code(doc, '''server.py（编排中心）
  ├── shared/encoding.py           # Windows GBK终端编码补丁
  ├── shared/config.py             # API密钥、路径、代理统一配置
  ├── shared/data_schema.py        # Pydantic模型 + 资料清单定义
  ├── phase1_parser.py             # Phase1: 多模态文档解析
  │     └── shared/mineru_client.py    # MinerU异步状态机客户端
  ├── phase2_analysis.py           # Phase2: 财务/科技/基本信息提取+收入验证
  │     └── shared/parsing.py          # 数字解析工具
  ├── phase2_calculator.py         # Phase2: 财务比率+杜邦分析+年度指标
  ├── phase2_inference.py          # Phase3.1: AI推理并行批量调度
  │     └── shared/llm_client.py       # Anthropic兼容客户端工厂
  ├── phase3_report.py             # Phase3: 报告生成入口
  │     └── docx_builder.py            # DOCX从零构建器（10章）
  └── phase4_committee.py          # Phase4: 5智能体×5轮贷审会辩论
        └── shared/llm_client.py''')

doc.add_page_break()

# ========================================================================
# 三、核心算法
# ========================================================================
add_h1(doc, '三、核心算法')

add_h2(doc, '3.1 财务报表智能提取算法')

add_body(doc,
    '财务报表提取是系统的数据基础。由于MinerU解析后的Markdown表格格式不完全标准'
    '（表名可能为中文长标题、列头可能有合并单元格痕迹、数字可能含千分位逗号），'
    '传统的按表名精确匹配和固定列位置取值方法在实际场景中频繁失效。本系统实现了三层自适应提取策略：')

add_h3(doc, '3.1.1 按内容匹配报表类型')

add_body(doc,
    '不依赖表名字符串匹配，而是扫描每张表格的第一列内容，按关键科目名定位报表类型：')

add_code(doc, '''def _find_table_by_content(tables, keywords):
    """
    遍历所有Markdown表格，扫描首列单元格文本，
    若包含目标关键词（如“资产总计”、“营业收入”）则匹配成功。
    返回匹配到的第一张表及其report_type标记。
    """
    for name, rows in tables.items():
        first_col = [row[0].strip() for row in rows if row]
        for kw_list, rtype in [
            (["资产总计", "总资产"], "balance_sheet"),
            (["营业收入", "营业总收入"], "income_statement"),
            (["经营活动", "经营性"], "cash_flow"),
        ]:
            if any(any(kw in cell for cell in first_col) for kw in kw_list):
                return name, rows, rtype
    return None, None, None''')

add_h3(doc, '3.1.2 按年份定位目标列')

add_body(doc,
    '不假设“最后一列即为最新数据”，而是解析表格列头中的年份信息（如“2025年”、“2024年度”），'
    '自动定位最近完整财年所在列，解决不同企业报表列排列不统一的问题：')

add_code(doc, '''def _find_target_column(header_row, prefer_year=None):
    """
    扫描表头行，搜索年份模式（2020-2026），
    返回最近完整财年的列索引。
    若当前为2026年5月，则最近完整财年为2025年。
    """
    year_cols = {}
    for i, cell in enumerate(header_row):
        match = re.search(r'(20[12]\\d)', str(cell))
        if match:
            year_cols[int(match.group(1))] = i
    if not year_cols:
        return -1  # 取最后一列作为回退
    target = prefer_year or _latest_complete_year()
    return year_cols.get(target, max(year_cols.keys()))''')

add_h3(doc, '3.1.3 合并/单体报表选择')

add_body(doc,
    '企业通常同时提供合并报表和单体（母公司）报表。系统优先使用合并报表（通过表名中的“合并”关键词识别），'
    '因为合并报表更能反映企业集团整体的财务状况，符合银行授信审批的惯例。若仅有一份报表，则直接使用。')

add_h2(doc, '3.2 收入交叉验证算法')

add_body(doc,
    '收入交叉验证是识别财务造假和经营异常的关键环节。系统从三个独立数据源提取收入数据，进行三角交叉比对：')

add_body(doc, '数据源一——财务报表：从利润表提取“营业收入”科目。')
add_body(doc,
    '数据源二——银行流水：通过正则匹配从银行流水中提取贷方发生额合计'
    '（匹配模式：贷方发生额合计、收入合计、本年累计流入等），取12个月累计值。')
add_body(doc,
    '数据源三——纳税申报表：通过正则匹配从税务申报表中提取申报营业收入'
    '（匹配模式：申报的营业收入、计税销售额等）。')

add_body(doc, '比对判定采用三级阈值：')

verify_headers = ['比对关系', '偏差率', '判定结论']
verify_rows = [
    ['报表营收 vs 银行流水', '≤10%', '基本吻合，收入真实性强'],
    ['报表营收 vs 银行流水', '10%-30%', '存在偏差，需进一步核查原因'],
    ['报表营收 vs 银行流水', '>30%', '重大偏差，收入真实性存疑'],
    ['报表营收 vs 纳税申报', '≤5%', '基本一致，税务合规'],
    ['报表营收 vs 纳税申报', '5%-15%', '存在差异，需核实原因'],
    ['报表营收 vs 纳税申报', '>15%', '重大差异，存在税务风险'],
]

add_table(doc, verify_headers, verify_rows, col_widths=[4.0, 2.5, 8.5])

add_h2(doc, '3.3 杜邦分析与财务指标体系')

add_body(doc,
    '系统基于国际通用的杜邦分析框架（DuPont Analysis），将净资产收益率（ROE）分解为三个驱动因子的乘积，'
    '帮助信贷员理解企业盈利能力的来源结构：')

add_body(doc, 'ROE = 销售净利率 × 资产周转率 × 权益乘数')

add_body(doc,
    '其中：销售净利率反映企业的成本控制和定价能力；资产周转率反映资产的运营效率；'
    '权益乘数（=总资产/所有者权益）反映财务杠杆水平。三个因子的乘积即为ROE。')

add_body(doc,
    '除杜邦分析外，系统还计算以下11项年度财务指标，覆盖盈利、偿债、营运三个维度：')

indicator_headers = ['维度', '指标名称', '计算公式']
indicator_rows = [
    ['盈利能力', '毛利率', '(营业收入 - 营业成本) / 营业收入 × 100%'],
    ['盈利能力', '销售利润率', '利润总额 / 营业收入 × 100%'],
    ['盈利能力', '净利率', '净利润 / 营业收入 × 100%'],
    ['偿债能力', '流动比率', '流动资产 / 流动负债'],
    ['偿债能力', '速动比率', '(流动资产 - 存货) / 流动负债'],
    ['偿债能力', '资产负债率', '总负债 / 总资产 × 100%'],
    ['偿债能力', '产权比率', '总负债 / 所有者权益'],
    ['偿债能力', '刚性负债净敞口', '（短期借款+应付票据+长期借款+应付债券）-（货币资金+交易性金融资产+应收票据）'],
    ['营运能力', '资产周转率', '营业收入 / 总资产'],
    ['营运能力', '应收账款周转天数', '360 / (营业收入 / 应收账款)'],
    ['营运能力', '存货周转天数', '360 / (营业成本 / 存货)'],
]

add_table(doc, indicator_headers, indicator_rows, col_widths=[2.0, 3.5, 9.5])

add_body(doc,
    '所有指标均支持多年对比计算。系统从Phase1的Markdown解析结果中按年份识别每列数据，'
    '对每个完整财年分别计算11项指标，在报告中以多年列表呈现并提供趋势分析。')

add_h2(doc, '3.4 多智能体辩论系统（贷审会模拟）')

add_body(doc,
    'Phase4的贷审会模拟是本系统在AI应用上的创新尝试。传统AI评估通常是单模型单次输出，'
    '难以模拟真实信贷审批中多方博弈、观点碰撞的复杂决策过程。本系统设计了一个结构化的多智能体辩论框架。')

add_h3(doc, '3.4.1 评委角色设计')

add_body(doc,
    '5位评委智能体分别代表贷审会中不同职能角色的思维方式和关注重点：')

committee_headers = ['评委', '角色定位', '核心关注点', '风格']
committee_rows = [
    ['牵头审批官\n（李行长）', '公司业务部总经理\n贷审会主持人', '综合风险收益、制度合规\n还款保障', '沉稳大气，善于倾听\n关键时刻一针见血'],
    ['风险审批官\n（王审）', '风险管理部资深审批官\n15年风控经验', '违约概率、担保充足性\n关联风险', '保守审慎，质疑乐观假设\n宁可错杀不可放过'],
    ['行业审批官\n（张博士）', '行业研究部总监\n产业经济学博士', '行业周期、竞争格局\n技术替代风险、政策环境', '战略思维，用数据说话\n关注赛道选择'],
    ['财务审批官\n（陈会计）', '授信审批部财务专家\n注册会计师', '盈利真实性、资产质量\n负债真实性、现金流健康度', '数据驱动，细节导向\n对财务造假信号极度敏感'],
    ['合规审批官\n（赵律师）', '法律合规部高级法务\n执业律师', '法律合规性、环保合规\n社会声誉、担保法律效力', '规则至上，零容忍\n不轻易为业务让步'],
]

add_table(doc, committee_headers, committee_rows, col_widths=[2.5, 3.5, 4.5, 4.5])

add_h3(doc, '3.4.2 辩论流程设计')

add_body(doc, '辩论分为5轮，从独立到对抗再到综合，逐步收敛：')

add_body(doc,
    '第1轮——独立审查：5位评委并行生成各自的初始立场陈述（互不通信），'
    '基于企业数据和自身角色标准给出初步的“支持/有条件支持/反对”判断及理由。')
add_body(doc,
    '第2轮——风险与合规交锋：主持人（李行长）引导风险审批官和合规审批官进行交锋讨论（6次发言），'
    '双方从各自视角质疑对方判断中的盲区。')
add_body(doc,
    '第3轮——行业与财务交锋：行业审批官与财务审批官进行对辩，'
    '从行业赛道和财务数据的交叉角度分析企业。')
add_body(doc,
    '第4轮——自由辩论：5位评委自由发表补充意见（7次发言），'
    '发言顺序按确定性优先级轮换（风险→行业→财务→合规→主持），不调用LLM做选择，减少API开销。')
add_body(doc,
    '第5轮——终审汇总：牵头审批官综合前4轮所有发言，输出结构化JSON终审结论，'
    '包括：授信结论（同意/有条件同意/否决）、核心理由、主要风险点、授信方案建议（金额/期限/担保条件）。')

add_h3(doc, '3.4.3 发言选择策略')

add_body(doc,
    '每轮辩论中的发言顺序采用确定性轮换而非LLM选择，以降低API开销并确保可复现。'
    '轮换基数为固定的优先级列表（风险→行业→财务→合规→主持），每轮从不同的起始位置开始循环。'
    '每位评委发言时能看到的上下文中包含本轮和之前轮次的所有辩论发言，确保信息充分传递。')

add_h2(doc, '3.5 AI推理并行批量调度')

add_body(doc,
    'Phase3.1的AI推理需要为42个定性字段生成专业分析文本。'
    '如果每个字段单独调用LLM API，按照每次调用约15秒计算，总耗时为42×15=630秒（约10分钟），'
    '严重影响用户体验。本系统设计了一套并行批量调度方案：')

add_h3(doc, '3.5.1 字段分组策略')

add_body(doc,
    '36个叙事型字段按语义关联性分为9个批次，每批2-7个字段在一个API请求中完成。'
    '另有3个结构化字段（前五大供应商、前五大客户、诉讼事件）需以JSON数组格式输出，独立调用。'
    '总计从42次串行调用压缩为12次API调用。')

add_h3(doc, '3.5.2 并发控制')

add_body(doc,
    '使用asyncio.Semaphore(3)控制最多3个批次的并发执行。该数值基于以下考量确定：'
    'Minimax API的速率限制（约5 QPS）、网络延迟的流水线效应、以及服务端内存占用。'
    '3并发是保持稳定性的安全上限。')

add_body(doc,
    '优化效果：API调用次数从42次降至12次（-71%），总耗时从约10分钟降至约90秒（-85%），'
    '且在推理进行中实时向用户展示批次完成进度，提升用户感知体验。')

doc.add_page_break()

# ========================================================================
# 四、关键技术选型与实现难点
# ========================================================================
add_h1(doc, '四、关键技术选型与实现难点')

add_h2(doc, '4.1 技术选型说明')

add_h3(doc, '4.1.1 LLM服务：Minimax（Anthropic兼容API）+ Claude 3.5 Sonnet')

add_body(doc, '选型考量：')
add_body(doc,
    '（1）国内可访问性：Google Gemini API在国内网络环境下访问不稳定，多次测试中出现超时和断连。'
    'Minimax提供的Anthropic兼容API端点（https://api.minimaxi.com/anthropic）国内访问延迟低、稳定性好。')
add_body(doc,
    '（2）模型能力：Claude 3.5 Sonnet在中文金融文本写作方面表现优异，生成的信贷分析报告语言专业、'
    '逻辑清晰，且严格遵循prompt中的数据约束（不编造数据、不确定时明确标注）。')
add_body(doc,
    '（3）SDK兼容性：使用官方anthropic-python SDK，仅需修改base_url即可无缝切换至Minimax端点，'
    '代码改动量极小（仅shared/llm_client.py中的客户端工厂函数）。')

add_h3(doc, '4.1.2 文档解析：MinerU云API（VLM）+ 本地pandas分流')

add_body(doc, '选型考量：')
add_body(doc,
    '（1）MinerU在PDF/图片的OCR和版面分析方面准确率高，能正确处理中文表格、印章遮挡、'
    '多栏布局等复杂场景，输出高质量的Markdown格式文本。')
add_body(doc,
    '（2）Excel分流是关键设计决策。将Excel文件交由pandas本地解析而非发送到MinerU云端，'
    '原因有二：一是财务报表的数字精确性至关重要，OCR引入的任何识别误差都可能传导至后续的'
    '财务比率计算，造成分析结论偏差；二是本地解析速度更快（秒级），不受云端排队影响。')
add_body(doc,
    '（3）MinerU客户端采用异步状态机设计，支持指数退避重试（1s/2s/4s/8s/16s）和递增间隔轮询（2→30秒），'
    '在保证成功率的同时避免对API的过度请求。')

add_h3(doc, '4.1.3 报告生成：python-docx从零构建（舍弃模板方案）')

add_body(doc, '选型考量：')
add_body(doc,
    'v4.x版本的“模板填空”方案存在根本性缺陷。Word模板基于XML Run层级存储文本，'
    '占位符可能被分散在多个Run元素中，导致文本替换不可靠；模板结构（341段+25表格）过于复杂，'
    '字段映射维护成本高；且模板中的预设格式难以灵活适配不同企业的差异化内容长度。')
add_body(doc,
    'v5.0完全舍弃Word模板，改用python-docx从零构建文档。优势包括：程序完全控制排版格式'
    '（字体、字号、行距、页边距）、根据数据类型动态选择渲染方式（文本段落、键值对、表格）、'
    '缺失数据统一以红色【待人工补充】标注、支持根据担保类型条件渲染章节内容。'
    '缺点是代码量较大（DocxReportBuilder约740行），但换来了完全的灵活性和可靠性。')

add_h2(doc, '4.2 实现难点与解决方案')

add_h3(doc, '难点1：财务报表表格的鲁棒匹配')

add_body(doc,
    '问题描述：不同企业的财务报表格式差异极大——表名可能是“资产负债表”、“合并资产负债表”、'
    '“资产负债表（续）”、“BALANCE SHEET”等；列排列可能按年（2023/2024/2025）或按期（期初/期末）；'
    '科目名可能含括号注释（如“流动资产合计（含递延）”）；部分表格可能是HTML格式嵌入Markdown。')

add_body(doc,
    '解决方案（v5.2）：采用三层自适应匹配策略。第一层：按内容而非按表名匹配——扫描表内第一列文本，'
    '检测是否包含“资产总计”、“营业收入”等特征科目名来判断报表类型。第二层：按年份定位目标列——'
    '解析表头中的年份数字，定位最近完整财年对应的列索引，而非假定最后一列即为最新数据。'
    '第三层：宽松科目名匹配——使用子串匹配（any(kw in label for kw in candidates)）替代精确匹配，'
    '适应括号注释、空格差异等变体。经过v4.6和v5.2两轮修复，报表提取的准确率从初始的约40%提升到90%以上。')

add_h3(doc, '难点2：LLM推理字段的批量调度与质量控制')

add_body(doc,
    '问题描述：42个推理字段如果逐一串行调用API，耗时长（约10分钟）且费用高。'
    '若简单地将所有字段合并为一个超大prompt，则单个字段的输出质量下降（LLM注意力分散），'
    '且单次请求token数过大增加超时风险。')

add_body(doc,
    '解决方案：采用分组+并发的双层调度策略。将36个叙事字段按所在报告章节的自然语义分为9组'
    '（如同一章的几个字段在一组），每组2-7个字段，共享一个batch prompt要求输出JSON。'
    '使用asyncio.Semaphore(3)控制并发度，在API限流和响应延迟之间取得平衡。'
    '每个batch prompt中明确要求每个字段200-400字的输出范围，并通过字段描述（description元数据）'
    '引导LLM的写作方向。最终效果：总API调用12次，总耗时约90秒，字段覆盖完整，'
    '质量稳定（因同一章节的字段共享上下文，分析逻辑更为一致）。')

add_h3(doc, '难点3：AI生成内容的幻觉风险控制')

add_body(doc,
    '问题描述：LLM在面对数据不足的字段时，可能“编造”看起来合理但实际未经确认的内容，'
    '在信贷审批场景中这种幻觉可能造成严重的错误决策。')

add_body(doc,
    '解决方案：在系统prompt中多层次约束：①明确要求“不编造任何未经确认的数据”；'
    '②要求所有财务分析结论都注明“根据所提供的财务报表推算”；③当现有资料不足以支撑分析时，'
    '强制使用固定话术“根据现有资料无法确定，建议实地调研补充”；'
    '④所有结构化/可计算指标（财务比率、增长率等）一律使用Python计算而非LLM推理，'
    '确保数字的精确可溯源；⑤人在环中的核对环节——信贷员可在Phase2完成后修正AI提取的数值。')

add_h3(doc, '难点4：HTML前端与FastAPI后端的会话状态同步')

add_body(doc,
    '问题描述：系统采用异步流水线，长耗时任务在后台执行，前端需要实时获取进度和中间结果。'
    '传统的请求-响应模式不能满足这种异步交互需求。')

add_body(doc,
    '解决方案：采用基于UUID的会话管理+前端轮询架构。每个浏览器Tab加载时生成UUID存入localStorage，'
    '后续所有API请求携带该UUID。服务端在内存dict中维护所有SessionState对象，包含工作流状态、'
    '进度信息、临时文件路径等。前端每2秒轮询GET /api/status/{session_id}获取最新状态，'
    '根据状态码驱动UI面板切换。该方案相比WebSocket更简单可靠（无连接管理和重连逻辑），'
    '适合本系统的并发规模（预计单机10-20并发用户）。后续可升级为Server-Sent Events（SSE）以减少轮询开销。')

add_h3(doc, '难点5：多智能体辩论的一致性维护')

add_body(doc,
    '问题描述：5轮辩论中每位评委的发言都基于前序所有发言的上下文。随着辩论推进，'
    '上下文长度不断增长（可能超过100K tokens），如何在不丢失关键信息的前提下控制token消耗？'
    '如何避免评委在后续轮次中“忘记”自己最初的立场？')

add_body(doc,
    '解决方案：①结构化的历史摘要——每轮辩论结束后，主持人（牵头审批官）将该轮核心观点压缩为'
    '摘要（而非直接拼接全部原始发言），下一轮辩论时将摘要作为上下文，既保留关键信息又控制长度。'
    '②立场追踪——每位评委的第1轮独立评审结论（初始立场）在所有后续轮次中始终作为上下文的一部分传递，'
    '确保评委不会“遗忘”初始判断。系统还会对比最终立场和初始立场，标注立场变化。'
    '③发言轮换机制——发言顺序按确定性优先级而非LLM自由选择，既降低API开销又确保每轮讨论覆盖所有视角。')

doc.add_page_break()

# ========================================================================
# 五、总结与展望
# ========================================================================
add_h1(doc, '五、总结与展望')

add_body(doc,
    '本项目实现了一套完整的银行对公信贷流程智能化辅助系统，覆盖了从多模态资料解析、'
    '财务数据提取与指标计算、AI辅助报告撰写，到多智能体贷审会模拟的全流程。'
    '系统已在真实企业资料（涵盖营业执照、财务报表、银行流水、纳税申报表、公司章程、专利证书等十余类文件）'
    '上完成端到端测试验证。')

add_body(doc, '主要技术成果：')

add_body(doc,
    '（1）多模态文档解析方面：通过MinerU VLM + pandas本地分流的混合架构，实现对6种以上格式的'
    '企业资料的自动解析，Excel财务数据确保数字精确性。')
add_body(doc,
    '（2）财务分析方面：实现了15+科目自动提取、11项指标/年的自动计算、杜邦分解和三方收入交叉验证，'
    '财务报表匹配准确率从初始40%提升至90%以上（经过v4.6和v5.2两轮核心算法改进）。')
add_body(doc,
    '（3）AI报告生成方面：通过分组并行的批量LLM调度策略，将42个推理字段的API调用从42次压缩至12次，'
    '总耗时从约10分钟降至约90秒（-85%），同时通过多层prompt约束控制AI幻觉风险。')
add_body(doc,
    '（4）多智能体辩论方面：构建了5角色×5轮的结构化辩论框架，输出包含结论、风险点和授信方案的'
    '结构化终审意见，为信贷决策提供多维度参考视角。')

add_body(doc, '未来优化方向：')
add_body(doc,
    '• 引入RAG（检索增强生成）机制，将银行内部的历史审批案例和行业基准数据纳入AI推理的知识库，'
    '提升分析的专业性和一致性。')
add_body(doc,
    '• 信贷员反馈闭环：将人工修正的数据作为微调样本，持续优化报表提取和AI推理的准确率。')
add_body(doc,
    '• 前端升级WebSocket或SSE实时推送，替代轮询机制，降低服务端负载并提升用户体验。')
add_body(doc,
    '• 引入更细粒度的风险评级模型（如PD/LGD违约概率模型），将AI定性分析与量化风控模型结合。')
add_body(doc,
    '• 支持更多LLM提供商的热切换（如Qwen、DeepSeek等国产模型），降低对单一供应商的依赖。')

# ── 保存 ──
output_path = 'D:/MBA/AI课/对公信贷流程智能化任务/技术说明文档_银行对公信贷智能化辅助系统.docx'
doc.save(output_path)
print('文档已生成: ' + output_path)
